import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import os

# 创建 Word 文档
doc = Document()

# 渲染公式为图片并保存
def render_latex_to_image(latex, filename):
    """渲染 LaTeX 公式为图片，并去除留白"""
    plt.figure(figsize=(1.5, 0.8))  # 调整图片宽高，减小默认大小
    plt.text(0.5, 0.5, f"${latex}$", fontsize=20, ha="center", va="center")  # 调整字体大小
    plt.axis("off")  # 移除坐标轴
    plt.tight_layout(pad=0)  # 去除留白
    plt.savefig(filename, bbox_inches="tight", dpi=300)  # 高分辨率保存
    plt.close()

# 从文本文件读取公式
file_path_txt = "latex_formulas.txt"
with open(file_path_txt, "r", encoding="utf-8") as file:
    latex_problems = file.readlines()

# 渲染每个公式并插入 Word 文档
for idx, formula in enumerate(latex_problems):
    formula = formula.strip()
    if formula:
        image_filename = f"formula_{idx + 1}.png"
        render_latex_to_image(formula, image_filename)  # 渲染图片
        doc.add_picture(image_filename, width=Inches(2))  # 插入图片，调整宽度为 2 英寸
        os.remove(image_filename)  # 删除生成的图片文件
        doc.add_paragraph("")  # 添加换行


# 保存 Word 文档
output_docx = "latex_formulas_word_with_images.docx"
doc.save(output_docx)

print(f"Word document saved at: {output_docx}")